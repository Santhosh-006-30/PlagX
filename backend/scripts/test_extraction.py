import io
import pypdf
import docx

def test_pdf():
    print("Testing PDF extraction...")
    try:
        # Create a simple PDF in memory
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.drawString(100, 750, "Hello PlagX PDF")
        c.save()
        content = buf.getvalue()
        
        pdf_reader = pypdf.PdfReader(io.BytesIO(content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        print(f"Extracted: {text.strip()}")
    except Exception as e:
        print(f"PDF extraction failed: {e}")

def test_docx():
    print("\nTesting DOCX extraction...")
    try:
        doc = docx.Document()
        doc.add_paragraph("Hello PlagX DOCX")
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()
        
        doc = docx.Document(io.BytesIO(content))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        print(f"Extracted: {text.strip()}")
    except Exception as e:
        print(f"DOCX extraction failed: {e}")

if __name__ == "__main__":
    test_pdf()
    test_docx()
